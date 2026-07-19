#!/usr/bin/env python3
"""Эталонный рендерер md -> docx для ВКР (vkr-writer skill).

Реализует все требования из SKILL.md, Фаза 4.5:
  * пользовательские стили VKR-H1/H2/H3 с outline_level,
    pageBreakBefore на H1;
  * красная строка 1.25 см только на body-абзацы (VKR-Body);
  * first_line_indent=0 на заголовки, captions, ячейки, библио;
  * hanging indent для библиографии (left=0, first_line=-0.75 см);
  * таблицы с обязательными границами (6 шт.), pct=5000;
  * парсер потребляет **Таблица N — ...** ровно один раз;
  * ![alt](path) -> картинка с одной подписью;
  * очистка ** / * / # / --- из текста;
  * TOC-поле \\o "1-3" \\h \\z \\u;
  * титульник — каждая строка отдельным абзацем (без \\n в run);
  * pre-save прогон docx_linter.lint_document().

Usage:
    python md_to_docx.py <input.md> <output.docx> [--meta meta.json]

Формат meta.json:
{
  "degree": "master",
  "topic": "...",
  "student": "Фамилия И. О.",
  "supervisor": "Фамилия И. О.",
  "specialty_code": "XX.XX.XX",
  "specialty_name": "...",
  "profile": "...",
  "year": "2026",
  "city": "г. Гатчина",
  "university_lines": ["АНО ВО «ЛАДИРО»", ...]
}
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


# ---------- константы форматирования ----------

FONT_NAME = "Times New Roman"
FONT_BODY = Pt(14)
FONT_TABLE = Pt(12)
LINE_SPACING_BODY = 1.5
LINE_SPACING_TABLE = 1.15
INDENT_BODY = Cm(1.25)
INDENT_BIBLIO_HANGING = Cm(0.75)  # first_line = -0.75, left = 0.75
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(1)
DEFAULT_IMAGE_WIDTH = Cm(16)


# ---------- XML helpers ----------


def _add_page_break_before_to_style(style) -> None:
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style.element.append(pPr)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pPr.append(OxmlElement("w:pageBreakBefore"))


def _set_outline_level(style, level_zero_based: int) -> None:
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style.element.append(pPr)
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(level_zero_based))


def _force_rfonts(run_or_style_element, name: str = FONT_NAME) -> None:
    rPr = run_or_style_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_or_style_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'w:ascii="{name}" w:hAnsi="{name}" '
            f'w:eastAsia="{name}" w:cs="{name}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(key), name)


def _set_cell_margins(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '<w:top w:w="55" w:type="dxa"/>'
        '<w:left w:w="85" w:type="dxa"/>'
        '<w:bottom w:w="55" w:type="dxa"/>'
        '<w:right w:w="85" w:type="dxa"/>'
        "</w:tcMar>"
    )
    tcPr.append(tcMar)


def _set_cell_shading(cell, color: str = "F2F2F2") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shd)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def _set_table_width_pct(table, pct: int = 5000) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{pct}" w:type="pct"/>')
    tblPr.append(tblW)


# ---------- стили ----------


def setup_styles(doc) -> None:
    # Normal: Times New Roman 14 pt, line 1.5, first_line=0 (на Normal индент НЕ вешаем!)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_BODY
    _force_rfonts(normal.element)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # VKR-Body: основной body-стиль с красной строкой 1.25 см
    body = _get_or_create_style(doc, "VKR-Body")
    body.base_style = normal
    body.font.name = FONT_NAME
    body.font.size = FONT_BODY
    pf = body.paragraph_format
    pf.first_line_indent = INDENT_BODY
    pf.line_spacing = LINE_SPACING_BODY
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _force_rfonts(body.element)

    # VKR-H1: 14 pt bold caps, CENTER, outline=0, pageBreakBefore
    h1 = _get_or_create_style(doc, "VKR-H1")
    h1.font.name = FONT_NAME
    h1.font.size = FONT_BODY
    h1.font.bold = True
    h1.font.all_caps = True
    pf = h1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    _set_outline_level(h1, 0)
    _add_page_break_before_to_style(h1)
    _force_rfonts(h1.element)

    # VKR-H2: 14 pt bold, CENTER, outline=1
    h2 = _get_or_create_style(doc, "VKR-H2")
    h2.font.name = FONT_NAME
    h2.font.size = FONT_BODY
    h2.font.bold = True
    pf = h2.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    _set_outline_level(h2, 1)
    _force_rfonts(h2.element)

    # VKR-H3: 14 pt bold italic, LEFT, outline=2
    h3 = _get_or_create_style(doc, "VKR-H3")
    h3.font.name = FONT_NAME
    h3.font.size = FONT_BODY
    h3.font.bold = True
    h3.font.italic = True
    pf = h3.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = INDENT_BODY
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    _set_outline_level(h3, 2)
    _force_rfonts(h3.element)

    # VKR-Caption: по левому/центру, без жирного, first_line=0
    cap = _get_or_create_style(doc, "VKR-Caption")
    cap.font.name = FONT_NAME
    cap.font.size = FONT_BODY
    pf = cap.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING_BODY
    pf.keep_with_next = True
    _force_rfonts(cap.element)

    # VKR-Biblio: hanging indent
    biblio = _get_or_create_style(doc, "VKR-Biblio")
    biblio.font.name = FONT_NAME
    biblio.font.size = FONT_BODY
    pf = biblio.paragraph_format
    pf.first_line_indent = -INDENT_BIBLIO_HANGING
    pf.left_indent = INDENT_BIBLIO_HANGING
    pf.line_spacing = LINE_SPACING_BODY
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _force_rfonts(biblio.element)

    # VKR-Cell: для абзацев внутри ячеек (first_line=0, line=1.15)
    cell_style = _get_or_create_style(doc, "VKR-Cell")
    cell_style.font.name = FONT_NAME
    cell_style.font.size = FONT_TABLE
    pf = cell_style.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.line_spacing = LINE_SPACING_TABLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _force_rfonts(cell_style.element)


def _get_or_create_style(doc, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


# ---------- настройка страницы ----------


def setup_page(doc) -> None:
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = MARGIN_LEFT
        sec.right_margin = MARGIN_RIGHT
        sec.top_margin = MARGIN_TOP
        sec.bottom_margin = MARGIN_BOTTOM


def set_document_language(doc, lang_ru: str = "ru-RU") -> None:
    """Устанавливает русский язык документа по умолчанию.

    Прописывает <w:lang w:val="ru-RU" .../> в:
    * <w:docDefaults>/<w:rPrDefault>/<w:rPr>  — для всего документа;
    * стилях Normal, VKR-H1/H2/H3, VKR-Body, VKR-Cell, VKR-Biblio, VKR-Caption.

    Без этого Word / LibreOffice может считать текст английским
    и подчёркивать всё красным (нет русского в словаре).
    """
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.insert(0, rPrDefault)
    default_rPr = rPrDefault.find(qn("w:rPr"))
    if default_rPr is None:
        default_rPr = OxmlElement("w:rPr")
        rPrDefault.append(default_rPr)
    old = default_rPr.find(qn("w:lang"))
    if old is not None:
        default_rPr.remove(old)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), lang_ru)
    lang.set(qn("w:eastAsia"), lang_ru)
    lang.set(qn("w:bidi"), "ar-SA")
    default_rPr.append(lang)

    for name in ("Normal", "VKR-H1", "VKR-H2", "VKR-H3", "VKR-Body",
                 "VKR-Cell", "VKR-Biblio", "VKR-Caption"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.insert(0, rPr)
        old = rPr.find(qn("w:lang"))
        if old is not None:
            rPr.remove(old)
        style_lang = OxmlElement("w:lang")
        style_lang.set(qn("w:val"), lang_ru)
        style_lang.set(qn("w:eastAsia"), lang_ru)
        style_lang.set(qn("w:bidi"), "ar-SA")
        rPr.append(style_lang)


def setup_page_numbers(doc) -> None:
    """Номера страниц в правом верхнем углу со 2-й страницы."""
    for sec in doc.sections:
        sectPr = sec._sectPr
        if sectPr.find(qn("w:titlePg")) is None:
            sectPr.append(OxmlElement("w:titlePg"))

        header = sec.header
        # clear existing header paragraphs
        for p in list(header.paragraphs):
            p.clear()
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE \\* MERGEFORMAT"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(placeholder)
        r.append(fld_end)


# ---------- базовые помощники для параграфов ----------


def _add_paragraph(doc, text: str, style: str, **run_kwargs):
    p = doc.add_paragraph(style=doc.styles[style])
    if text:
        _append_inline(p, text, **run_kwargs)
    return p


def _append_inline(p, text: str, *, force_bold: bool = False,
                   force_italic: bool = False, size: Pt = None) -> None:
    """Разбить text на runs, обрабатывая **жирный** и *курсив*."""
    if not text:
        return
    # clear existing empty run if present
    for segment in re.split(r"(\*\*.*?\*\*|\*[^*\n]+\*)", text):
        if not segment:
            continue
        bold = force_bold
        italic = force_italic
        if segment.startswith("**") and segment.endswith("**") and len(segment) >= 4:
            segment = segment[2:-2]
            bold = True
        elif (segment.startswith("*") and segment.endswith("*")
              and len(segment) >= 2):
            segment = segment[1:-1]
            italic = True
        run = p.add_run(segment)
        run.font.name = FONT_NAME
        if size is not None:
            run.font.size = size
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        _force_rfonts(run._r)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- титульник ----------


def add_title_page(doc, meta: dict) -> None:
    def line(text, *, bold=False, size=FONT_BODY, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(0)):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.space_after = space_after
        pf.line_spacing = LINE_SPACING_BODY
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size
        if bold:
            run.font.bold = True
        _force_rfonts(run._r)
        return p

    def spacer(n: int = 1):
        for _ in range(n):
            line("", space_after=Pt(0))

    # Шапка университета — каждая строка отдельным абзацем
    uni_lines = meta.get("university_lines", [
        "Автономное образовательное учреждение",
        "высшего образования Ленинградской области",
        "«Государственный институт экономики, финансов, права и технологий»",
    ])
    spacer(1)
    for text in uni_lines:
        line(text)

    spacer(1)
    line(meta.get("faculty", "Факультет экономический"), bold=True,
         space_after=Pt(18))

    line("Допустить к защите", align=WD_ALIGN_PARAGRAPH.RIGHT)
    line("Зав. кафедрой ___________________", align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=Pt(24))

    degree_label = ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА МАГИСТРА"
                    if meta.get("degree") == "master"
                    else "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА БАКАЛАВРА")
    line(degree_label, bold=True, size=Pt(16), space_after=Pt(12))
    line(f"на тему «{meta.get('topic', '…')}»", bold=True, space_after=Pt(18))

    spec = f"{meta.get('specialty_code', '')} {meta.get('specialty_name', '')}".strip()
    if spec:
        line(f"Направление подготовки {spec}")
    if meta.get("profile"):
        line(f"Профиль — {meta['profile']}", space_after=Pt(18))

    line("Выполнил(а):")
    line(meta.get("student", "_______________________"), space_after=Pt(12))
    line("Научный руководитель:")
    line(meta.get("supervisor", "_______________________"))

    spacer(3)
    line(meta.get("city", "г. Гатчина"))
    line(str(meta.get("year", "2026")))


# ---------- оглавление ----------


def add_toc(doc) -> None:
    _add_page_break(doc)
    _add_paragraph(doc, "СОДЕРЖАНИЕ", style="VKR-H1")
    p = doc.add_paragraph(style=doc.styles["Normal"])
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_BODY
    _force_rfonts(run._r)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Для обновления оглавления нажмите F9 или ПКМ → Обновить поле."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


# ---------- таблицы ----------


def _looks_numeric(text: str) -> bool:
    t = text.replace(" ", "").replace(",", ".").replace("—", "").replace("-", "").strip()
    if not t:
        return False
    return bool(re.match(r"^-?\d+(?:\.\d+)?%?$", t))


def _clean_md_inline(text: str) -> str:
    """Удалить ** и * из cell-текста, не ломая числа."""
    return re.sub(r"\*+", "", text)


def add_markdown_table(doc, md_lines: list[str]) -> None:
    """Построить docx-таблицу из markdown `| a | b |` строк."""
    rows: list[list[str]] = []
    for line in md_lines:
        s = line.strip()
        if not s:
            continue
        if re.match(r"^\|[-:\s|]+\|?$", s):
            continue
        cells = [c.strip() for c in s.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    if len(rows) < 2:
        return

    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    _set_table_width_pct(table, 5000)

    # шапка (w:tblHeader на первой строке для повтора на новой странице)
    header_row = table.rows[0]
    header_row_trPr = header_row._tr.get_or_add_trPr()
    if header_row_trPr.find(qn("w:tblHeader")) is None:
        header_row_trPr.append(OxmlElement("w:tblHeader"))

    for ri, row_data in enumerate(rows):
        for ci, raw_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.style = doc.styles["VKR-Cell"]
            pf = p.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.line_spacing = LINE_SPACING_TABLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            txt = _clean_md_inline(raw_text)
            if ri == 0:
                _set_cell_shading(cell, "F2F2F2")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(txt)
                run.font.name = FONT_NAME
                run.font.size = FONT_TABLE
                run.font.bold = True
                _force_rfonts(run._r)
            else:
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                               if _looks_numeric(txt)
                               else WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(txt)
                run.font.name = FONT_NAME
                run.font.size = FONT_TABLE
                _force_rfonts(run._r)


# ---------- изображения и подписи ----------


def add_image_with_caption(doc, image_path: Path, caption_text: str | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=DEFAULT_IMAGE_WIDTH)
    except Exception as exc:
        run.text = f"[Рисунок недоступен: {image_path.name} — {exc}]"
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.size = Pt(12)
        run.font.name = FONT_NAME
        _force_rfonts(run._r)
    if caption_text:
        cap = _add_paragraph(doc, caption_text, style="VKR-Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_caption(doc, caption_text: str) -> None:
    p = _add_paragraph(doc, caption_text, style="VKR-Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_source_note(doc, text: str) -> None:
    p = doc.add_paragraph(style=doc.styles["VKR-Caption"])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.italic = True
    _force_rfonts(run._r)


def add_bibliography_entry(doc, text: str) -> None:
    p = doc.add_paragraph(style=doc.styles["VKR-Biblio"])
    _append_inline(p, text)


def add_list_item(doc, text: str, *, ordered: bool = False, number: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = INDENT_BODY
    pf.line_spacing = LINE_SPACING_BODY
    prefix = f"{number}. " if ordered else "— "
    run = p.add_run(prefix)
    run.font.name = FONT_NAME
    run.font.size = FONT_BODY
    _force_rfonts(run._r)
    _append_inline(p, text)


# ---------- ПАРСЕР ----------


IMAGE_RE = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")
TABLE_CAPTION_RE = re.compile(r"^\*\*(Таблица\s+[\dА-Яа-я\.]+\s*[—\-–].*?)\*\*\s*$")
FIGURE_CAPTION_RE = re.compile(r"^\*\*(Рисунок\s+[\dА-Яа-я\.]+\s*[—\-–].*?)\*\*\s*$")
CHAPTER_PREFIXES = ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК", "СОДЕРЖАНИЕ",
                    "ПРИЛОЖЕНИЕ", "ПРИЛОЖЕНИЯ")


@dataclass
class ParserState:
    in_bibliography: bool = False
    in_appendix: bool = False
    pending_table_caption: str | None = None
    pending_figure_caption: str | None = None
    base_dir: Path = field(default_factory=Path.cwd)


def _is_chapter_like(title: str) -> bool:
    upper = title.strip().upper()
    if any(upper.startswith(pref) for pref in CHAPTER_PREFIXES):
        return True
    if re.match(r"^ГЛАВА\s+\d+", title, re.IGNORECASE):
        return True
    if re.match(r"^\d+\s", title.strip()):
        return True
    return False


def _strip_chapter_prefix(title: str) -> str:
    m = re.match(r"^ГЛАВА\s+(\d+)\s*\.?\s*(.*)$", title, re.IGNORECASE)
    if m:
        return f"ГЛАВА {m.group(1)}. {m.group(2).strip()}".strip()
    return title


def parse_markdown(doc, md_text: str, base_dir: Path) -> None:
    lines = md_text.split("\n")
    state = ParserState(base_dir=base_dir)
    i = 0

    def flush_pending_figure_caption():
        if state.pending_figure_caption:
            cap = _add_paragraph(doc, state.pending_figure_caption,
                                 style="VKR-Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            state.pending_figure_caption = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # empty line
        if not stripped:
            i += 1
            continue

        # --- horizontal rule — СКИПАЕМ
        if stripped == "---" or re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # H1 — # Заголовок
        if re.match(r"^#\s+", line):
            title = _strip_chapter_prefix(re.sub(r"^#\s+", "", line).strip())
            # Определяем секции
            upper = title.upper()
            state.in_bibliography = upper.startswith("СПИСОК")
            state.in_appendix = upper.startswith("ПРИЛОЖЕНИ")
            _add_paragraph(doc, title, style="VKR-H1")
            i += 1
            continue

        # H2
        if re.match(r"^##\s+", line):
            title = re.sub(r"^##\s+", "", line).strip()
            title = re.sub(r"^(\d+(?:\.\d+)*)\.(\s)", r"\1\2", title)
            # H2, начинающийся с "Приложение X" — трактуем как H1
            # (отдельная страница, стиль VKR-H1)
            if re.match(r"^Приложение\s+[А-ЯA-Z]", title):
                _add_paragraph(doc, title, style="VKR-H1")
            else:
                _add_paragraph(doc, title, style="VKR-H2")
            i += 1
            continue

        # H3
        if re.match(r"^###\s+", line):
            title = re.sub(r"^###\s+", "", line).strip()
            title = re.sub(r"^(\d+(?:\.\d+)*)\.(\s)", r"\1\2", title)
            _add_paragraph(doc, title, style="VKR-H3")
            i += 1
            continue

        # Table caption **Таблица N — ...**
        m_tcap = TABLE_CAPTION_RE.match(stripped)
        if m_tcap:
            state.pending_table_caption = m_tcap.group(1).strip()
            i += 1
            continue

        # Figure caption **Рисунок N — ...**
        m_fcap = FIGURE_CAPTION_RE.match(stripped)
        if m_fcap:
            state.pending_figure_caption = m_fcap.group(1).strip()
            i += 1
            continue

        # Image ![alt](path)
        m_img = IMAGE_RE.match(stripped)
        if m_img:
            path = m_img.group(1)
            img_path = (state.base_dir / path) if not Path(path).is_absolute() \
                else Path(path)
            add_image_with_caption(doc, img_path, state.pending_figure_caption)
            state.pending_figure_caption = None
            i += 1
            continue

        # Markdown table
        if (stripped.startswith("|") and i + 1 < len(lines)
                and re.match(r"^\|[-:\s|]+\|?$", lines[i + 1].strip())):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # emit pending caption сначала
            if state.pending_table_caption:
                add_table_caption(doc, state.pending_table_caption)
                state.pending_table_caption = None
            add_markdown_table(doc, table_lines)
            continue

        # Источник: ...
        if stripped.startswith("Источник:"):
            add_source_note(doc, stripped)
            i += 1
            continue

        # Bibliography entry (внутри секции "СПИСОК ИСТОЧНИКОВ")
        if state.in_bibliography and re.match(r"^\d+\.\s+", stripped):
            add_bibliography_entry(doc, stripped)
            i += 1
            continue

        # Ordered list item (вне bibliography)
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num and not state.in_bibliography:
            add_list_item(doc, m_num.group(2), ordered=True,
                          number=m_num.group(1))
            i += 1
            continue

        # Bullet list
        if stripped.startswith(("- ", "* ", "— ", "– ")):
            text = re.sub(r"^[-*—–]\s+", "", stripped)
            add_list_item(doc, text, ordered=False)
            i += 1
            continue

        # обычный body
        # Если накопилась таблица-подпись но таблицы нет — всё равно выдадим caption
        if state.pending_table_caption:
            add_table_caption(doc, state.pending_table_caption)
            state.pending_table_caption = None
        # Figure caption без image — выдадим как подпись
        flush_pending_figure_caption()

        p = doc.add_paragraph(style=doc.styles["VKR-Body"])
        _append_inline(p, stripped)
        i += 1

    # hanging flushes
    if state.pending_table_caption:
        add_table_caption(doc, state.pending_table_caption)
    flush_pending_figure_caption()


# ---------- main build ----------


def build_docx(md_path: Path, docx_path: Path, meta: dict,
               base_dir: Path | None = None) -> None:
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    set_document_language(doc)

    # Титульник БЕЗ ведущего page break!
    add_title_page(doc, meta)
    # Оглавление с новой страницы
    add_toc(doc)
    # Body — парсер расставит # на VKR-H1 с pageBreakBefore
    md_text = md_path.read_text(encoding="utf-8")
    parse_markdown(doc, md_text, base_dir=base_dir or md_path.parent)

    setup_page_numbers(doc)

    # pre-save lint
    try:
        from docx_linter import lint_document  # type: ignore
    except ImportError:
        sys.path.insert(0, str(Path(__file__).parent))
        from docx_linter import lint_document  # type: ignore

    report = lint_document(doc)
    if report.critical:
        print("WARNING: docx_linter обнаружил критичные дефекты ДО сохранения:",
              file=sys.stderr)
        for issue in report.critical[:10]:
            print(f"  [{issue.code}] {issue.message}", file=sys.stderr)

    doc.save(str(docx_path))
    print(f"Saved {docx_path}")
    print(f"  critical: {len(report.critical)}  warnings: {len(report.warnings)}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="md -> docx рендерер для ВКР (vkr-writer skill)."
    )
    parser.add_argument("input_md", help="Входной markdown-файл")
    parser.add_argument("output_docx", help="Выходной docx")
    parser.add_argument("--meta", help="JSON с метаданными титульника")
    args = parser.parse_args(argv)

    md_path = Path(args.input_md)
    docx_path = Path(args.output_docx)
    if not md_path.exists():
        print(f"ERROR: input not found: {md_path}", file=sys.stderr)
        return 1

    meta = {}
    if args.meta:
        meta_path = Path(args.meta)
        if meta_path.exists():
            meta = json.loads(meta_path.read_text(encoding="utf-8"))

    build_docx(md_path, docx_path, meta)
    return 0


if __name__ == "__main__":
    sys.exit(main())
