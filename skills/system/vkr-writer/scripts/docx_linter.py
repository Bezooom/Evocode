#!/usr/bin/env python3
"""DOCX linter for vkr-writer skill.

Проверяет DOCX-файл ВКР по 6 категориям дефектов (см. SKILL.md, Фаза 4.5).
Выходной формат — JSON с issues/warnings + exit code:
  0 — всё хорошо,
  1 — критичные проблемы (category=critical),
  2 — только warnings.

Usage:
    python docx_linter.py <path-to.docx>            # печать JSON и exit
    python docx_linter.py <path-to.docx> --quiet    # только exit code
    python docx_linter.py <path-to.docx> --human    # человеко-читаемый отчёт
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx",
          file=sys.stderr)
    sys.exit(3)


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@dataclass
class Issue:
    category: str            # critical | warning
    code: str                # short id
    message: str
    location: str = ""       # para idx / table idx etc.


@dataclass
class Report:
    file: str
    issues: list[Issue] = field(default_factory=list)

    def add(self, category: str, code: str, message: str, location: str = "") -> None:
        self.issues.append(Issue(category, code, message, location))

    @property
    def critical(self) -> list[Issue]:
        return [i for i in self.issues if i.category == "critical"]

    @property
    def warnings(self) -> list[Issue]:
        return [i for i in self.issues if i.category == "warning"]

    def to_dict(self) -> dict:
        return {
            "file": self.file,
            "critical_count": len(self.critical),
            "warning_count": len(self.warnings),
            "issues": [asdict(i) for i in self.issues],
        }


# ---------- helpers ----------


def iter_all_paragraphs(doc) -> Iterable[tuple[str, int, object]]:
    """Yield (location, index, paragraph) for body + table cells."""
    for i, p in enumerate(doc.paragraphs):
        yield ("body", i, p)
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield (f"T{ti}[{ri},{ci}]", pi, p)


def has_page_break_in(p) -> bool:
    return any(
        br.get(qn("w:type")) == "page"
        for run in p.runs
        for br in run._element.iter(qn("w:br"))
    )


def has_page_break_before(p) -> bool:
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
        return True
    return False


# ---------- checks ----------


def check_first_page(doc, rep: Report) -> None:
    if not doc.paragraphs:
        rep.add("critical", "empty_doc", "Документ не содержит абзацев")
        return
    first = doc.paragraphs[0]
    if has_page_break_in(first) and not first.text.strip():
        rep.add(
            "critical",
            "leading_page_break",
            "Первый абзац — пустой с page break: титульник уезжает на 2-ю страницу",
            "para[0]",
        )


def check_markdown_artifacts(doc, rep: Report) -> None:
    bad = {
        "md_heading": re.compile(r"^#{1,6}\s"),
        "md_bold":    re.compile(r"\*\*[^*\s][^*]*\*\*"),
        "md_italic":  re.compile(r"(?<!\*)\*[^*\s][^*\n]*\*(?!\*)"),
        "md_hr":      re.compile(r"^-{3,}\s*$"),
        "md_bullet":  re.compile(r"^\s*[-*]\s+\w"),
        "md_image":   re.compile(r"!\[[^\]]*\]\([^)]+\)"),
        "md_link":    re.compile(r"(?<!!)\[[^\]]+\]\([^)]+\)"),
    }
    counters = {k: 0 for k in bad}
    samples = {k: [] for k in bad}

    for loc, idx, p in iter_all_paragraphs(doc):
        text = p.text
        for code, patt in bad.items():
            if patt.search(text):
                counters[code] += 1
                if len(samples[code]) < 3:
                    samples[code].append((loc, idx, text[:80]))

    for code, n in counters.items():
        if n == 0:
            continue
        first_samples = "; ".join(
            f"{loc}/{idx}:{txt!r}" for loc, idx, txt in samples[code]
        )
        rep.add(
            "critical",
            code,
            f"Markdown-артефакт '{code}' найден {n} раз(а). Примеры: {first_samples}",
        )


def check_styles(doc, rep: Report) -> None:
    total = len(doc.paragraphs)
    if total == 0:
        return
    normal = sum(1 for p in doc.paragraphs if p.style.name == "Normal")
    ratio = normal / total
    if ratio >= 0.9:
        rep.add(
            "critical",
            "all_normal_style",
            f"{normal}/{total} ({ratio:.0%}) абзацев имеют стиль 'Normal'. "
            "Заголовки должны использовать Heading 1/2/3 или VKR-H1/H2/H3.",
        )

    # Паттерны реальных chapter-headings верхнего уровня.
    # СОДЕРЖАНИЕ — спец-страница, в TOC не попадает, стиль может быть Normal,
    # поэтому из проверки исключается.
    chapter_re = re.compile(
        r"^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|СПИСОК\b|ПРИЛОЖЕНИЯ\b"
        r"|ГЛАВА\s+\d+|ПРИЛОЖЕНИЕ\s+[А-ЯA-Z](?:\b|\.))",
        re.UNICODE,
    )
    for i, p in enumerate(doc.paragraphs):
        up = p.text.strip().upper()
        if chapter_re.match(up):
            if p.style.name == "Normal":
                rep.add(
                    "critical",
                    "heading_wrong_style",
                    f"Заголовок {p.text.strip()[:40]!r} имеет стиль 'Normal', "
                    "должен быть Heading 1 / VKR-H1",
                    f"para[{i}]",
                )


def check_tables(doc, rep: Report) -> None:
    for ti, t in enumerate(doc.tables):
        tbl = t._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            rep.add("critical", "table_no_tblPr",
                    f"Таблица {ti}: отсутствует tblPr", f"T{ti}")
            continue
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            rep.add("critical", "table_no_borders",
                    f"Таблица {ti}: нет tblBorders — невидимая таблица",
                    f"T{ti}")
        else:
            required = {"top", "left", "bottom", "right", "insideH", "insideV"}
            present = {b.tag.split("}")[-1] for b in borders}
            missing = required - present
            if missing:
                rep.add("critical", "table_borders_incomplete",
                        f"Таблица {ti}: отсутствуют границы {sorted(missing)}",
                        f"T{ti}")
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            rep.add("warning", "table_no_width",
                    f"Таблица {ti}: не задана ширина tblW", f"T{ti}")
        else:
            w_type = tblW.get(qn("w:type"))
            if w_type == "auto":
                rep.add("warning", "table_width_auto",
                        f"Таблица {ti}: ширина type='auto' "
                        "— рекомендуется pct=5000",
                        f"T{ti}")

        # first_line_indent в ячейках должен быть 0
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    fli = p.paragraph_format.first_line_indent
                    if fli is not None and fli.cm > 0.1:
                        rep.add(
                            "critical",
                            "cell_first_line_indent",
                            f"Таблица {ti}[{ri},{ci}]: "
                            f"first_line_indent={fli.cm:.2f}см, должен быть 0",
                            f"T{ti}[{ri},{ci}]",
                        )
                        break


def check_duplicate_captions(doc, rep: Report) -> None:
    caption_re = re.compile(r"^\**\s*(Таблица|Рисунок)\s+[\dА-Яа-я\.]+")

    def canon(text: str) -> str:
        return re.sub(r"[*\s]+", " ", text).strip().lower()

    prev = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if caption_re.match(t):
            c = canon(t)
            if prev and c == prev[1]:
                rep.add(
                    "critical",
                    "duplicate_caption",
                    f"Дубликат подписи: para[{prev[0]}] и para[{i}]: "
                    f"{t[:60]!r}",
                    f"para[{i}]",
                )
            prev = (i, c)
        elif t:  # non-caption non-empty resets streak
            prev = None


def check_titlepage_multiline(doc, rep: Report) -> None:
    for i, p in enumerate(doc.paragraphs[:15]):
        if "\n" in p.text:
            rep.add(
                "critical",
                "titlepage_newline_in_run",
                f"para[{i}] содержит \\n внутри run — "
                "многострочная строка должна быть несколькими абзацами: "
                f"{p.text[:60]!r}",
                f"para[{i}]",
            )


def check_page_breaks_before_chapters(doc, rep: Report) -> None:
    chapter_re = re.compile(
        r"^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|СОДЕРЖАНИЕ|СПИСОК\b|ПРИЛОЖЕНИЯ\b"
        r"|ГЛАВА\s+\d+|ПРИЛОЖЕНИЕ\s+[А-ЯA-Z](?:\b|\.))",
        re.UNICODE,
    )
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        up = p.text.strip().upper()
        if not chapter_re.match(up):
            continue
        if i == 0:
            continue  # no break needed before first paragraph
        # check page break in self (inline), in pBB or in previous
        if has_page_break_in(p) or has_page_break_before(p):
            continue
        if i > 0 and has_page_break_in(paragraphs[i - 1]):
            continue
        # check style has pageBreakBefore
        try:
            style = p.style
            if style is not None:
                style_pPr = style.element.find(qn("w:pPr"))
                if style_pPr is not None and \
                        style_pPr.find(qn("w:pageBreakBefore")) is not None:
                    continue
        except Exception:
            pass
        rep.add(
            "warning",
            "no_page_break_before_chapter",
            f"Нет page break перед para[{i}] {p.text.strip()[:50]!r}",
            f"para[{i}]",
        )


def _effective_first_line_indent(p):
    """Возвращает first_line_indent с учётом наследования от стиля."""
    direct = p.paragraph_format.first_line_indent
    if direct is not None:
        return direct
    style = p.style
    while style is not None:
        pf = style.paragraph_format
        if pf.first_line_indent is not None:
            return pf.first_line_indent
        style = getattr(style, "base_style", None)
    return None


def _effective_left_indent(p):
    direct = p.paragraph_format.left_indent
    if direct is not None:
        return direct
    style = p.style
    while style is not None:
        pf = style.paragraph_format
        if pf.left_indent is not None:
            return pf.left_indent
        style = getattr(style, "base_style", None)
    return None


def check_body_first_line_indent(doc, rep: Report) -> None:
    body = []
    for p in doc.paragraphs:
        if p.style.name in ("Normal", "VKR-Body") and len(p.text) > 100:
            body.append(p)
    if not body:
        return
    with_indent = 0
    for p in body:
        fli = _effective_first_line_indent(p)
        if fli is not None and fli.cm > 0.5:
            with_indent += 1
    ratio = with_indent / len(body)
    if ratio < 0.5:
        rep.add(
            "warning",
            "no_first_line_indent",
            f"Только {with_indent}/{len(body)} ({ratio:.0%}) длинных "
            f"body-абзацев имеют красную строку (1.25 см). Требование ГОСТ.",
        )


def check_bibliography_hanging(doc, rep: Report) -> None:
    biblio_re = re.compile(r"^\d+\.\s+[А-ЯA-Z]")
    in_biblio = False
    bad = 0
    total = 0
    for p in doc.paragraphs:
        up = p.text.strip().upper()
        if up.startswith("СПИСОК"):
            in_biblio = True
            continue
        if up.startswith("ПРИЛОЖЕНИ"):
            in_biblio = False
            continue
        if in_biblio and biblio_re.match(p.text.strip()):
            total += 1
            fli = _effective_first_line_indent(p)
            li = _effective_left_indent(p)
            if fli is None or fli.cm >= 0 or li is None:
                bad += 1
    if total > 0 and bad / total > 0.5:
        rep.add(
            "warning",
            "biblio_no_hanging_indent",
            f"Библиография: {bad}/{total} записей без висячей строки "
            "(ГОСТ Р 7.0.100-2018 требует hanging indent 0.75–1.25 см).",
        )


def check_toc_field(doc, rep: Report) -> None:
    for p in doc.paragraphs[:40]:
        for elem in p._element.iter(qn("w:instrText")):
            if elem.text and "TOC" in elem.text.upper():
                return
    rep.add(
        "warning",
        "no_toc_field",
        "Оглавление не оформлено как поле TOC — собрано вручную "
        "(без номеров страниц).",
    )


def check_document_language(doc, rep: Report) -> None:
    """Проверяет, что язык документа по умолчанию — русский.

    Ищет <w:lang w:val="ru-*"/> в <w:docDefaults>/<w:rPrDefault>/<w:rPr>
    или в стиле Normal. Если нигде не найдено — warning: Word будет
    подчёркивать весь кириллический текст красным.
    """
    styles_element = doc.styles.element

    # docDefaults
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is not None:
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn("w:rPr"))
            if rPr is not None:
                lang = rPr.find(qn("w:lang"))
                if lang is not None:
                    val = lang.get(qn("w:val")) or ""
                    if val.lower().startswith("ru"):
                        return

    # fallback: стиль Normal
    try:
        normal = doc.styles["Normal"]
        rPr = normal.element.find(qn("w:rPr"))
        if rPr is not None:
            lang = rPr.find(qn("w:lang"))
            if lang is not None:
                val = lang.get(qn("w:val")) or ""
                if val.lower().startswith("ru"):
                    return
    except KeyError:
        pass

    rep.add(
        "warning",
        "document_language_not_ru",
        "Язык документа по умолчанию не установлен как русский "
        "(w:lang w:val=\"ru-RU\"). Word будет подчёркивать текст красным, "
        "проверка орфографии и переносы сломаются.",
    )


def check_margins(doc, rep: Report) -> None:
    for si, sec in enumerate(doc.sections):
        pairs = [
            ("left", sec.left_margin.cm, 3.0, 0.2),
            ("right", sec.right_margin.cm, 1.0, 0.3),
            ("top", sec.top_margin.cm, 2.0, 0.2),
            ("bottom", sec.bottom_margin.cm, 2.0, 0.2),
        ]
        for name, val, expected, tol in pairs:
            if abs(val - expected) > tol:
                rep.add(
                    "warning",
                    f"margin_{name}",
                    f"Секция {si}: {name} margin = {val:.2f}см "
                    f"(ожидается {expected}см ±{tol})",
                )


# ---------- main ----------


def lint_doc(doc_path: Path) -> Report:
    rep = Report(file=str(doc_path))
    doc = Document(str(doc_path))
    check_first_page(doc, rep)
    check_markdown_artifacts(doc, rep)
    check_styles(doc, rep)
    check_tables(doc, rep)
    check_duplicate_captions(doc, rep)
    check_titlepage_multiline(doc, rep)
    check_page_breaks_before_chapters(doc, rep)
    check_body_first_line_indent(doc, rep)
    check_bibliography_hanging(doc, rep)
    check_toc_field(doc, rep)
    check_margins(doc, rep)
    check_document_language(doc, rep)
    return rep


def lint_document(doc) -> Report:
    """Entry point для вызова на уже открытом Document (из генератора)."""
    rep = Report(file="<in-memory>")
    check_first_page(doc, rep)
    check_markdown_artifacts(doc, rep)
    check_styles(doc, rep)
    check_tables(doc, rep)
    check_duplicate_captions(doc, rep)
    check_titlepage_multiline(doc, rep)
    check_page_breaks_before_chapters(doc, rep)
    check_body_first_line_indent(doc, rep)
    check_bibliography_hanging(doc, rep)
    check_toc_field(doc, rep)
    check_margins(doc, rep)
    check_document_language(doc, rep)
    return rep


def format_human(rep: Report) -> str:
    lines = [f"DOCX linter report — {rep.file}"]
    lines.append(f"  critical: {len(rep.critical)}   warnings: {len(rep.warnings)}")
    lines.append("")
    if rep.critical:
        lines.append("CRITICAL ISSUES:")
        for i in rep.critical:
            lines.append(f"  [{i.code}] {i.message}")
            if i.location:
                lines.append(f"      at {i.location}")
        lines.append("")
    if rep.warnings:
        lines.append("WARNINGS:")
        for i in rep.warnings:
            lines.append(f"  [{i.code}] {i.message}")
            if i.location:
                lines.append(f"      at {i.location}")
        lines.append("")
    if not rep.issues:
        lines.append("OK — no issues.")
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Linter DOCX-файла ВКР (см. SKILL.md vkr-writer)."
    )
    parser.add_argument("docx", help="Путь к DOCX-файлу")
    parser.add_argument("--quiet", action="store_true",
                        help="Не печатать отчёт, только exit code")
    parser.add_argument("--human", action="store_true",
                        help="Человеко-читаемый отчёт вместо JSON")
    args = parser.parse_args(argv)

    path = Path(args.docx)
    if not path.exists():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 3

    rep = lint_doc(path)

    if not args.quiet:
        if args.human:
            print(format_human(rep))
        else:
            print(json.dumps(rep.to_dict(), ensure_ascii=False, indent=2))

    if rep.critical:
        return 1
    if rep.warnings:
        return 2
    return 0


if __name__ == "__main__":
    sys.exit(main())
