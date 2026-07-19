#!/usr/bin/env python3
"""Рендерер md -> docx для ДПО-программ (dpo-program skill).

Реализует генерацию docx из markdown-файлов тем ДПО с ГОСТ-оформлением.
Поддерживает:
  * Титульный лист ДПО-программы
  * Стили заголовков H1/H2/H3 с outline_level
  * Красная строка 1.25 см для body
  * Таблицы с обязательными границами
  * Списки (нумерованные и маркированные)
  * Markdown inline (**bold**, *italic*)
  * Автогенерация оглавления (TOC-поле)

Usage:
    python dpo_md_to_docx.py --meta meta.json --topics topic-1.md topic-3.md --output output.docx
    python dpo_md_to_docx.py --meta meta.json --all --output output.docx
"""

from __future__ import annotations

import argparse
import glob
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

# ---------- константы ГОСТ ----------

FONT_NAME = "Times New Roman"
FONT_BODY = Pt(14)
FONT_TABLE = Pt(12)
LINE_SPACING_BODY = 1.5
LINE_SPACING_TABLE = 1.15
INDENT_BODY = Cm(1.25)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(1.5)

# ---------- XML helpers ----------


def _rf(element, tag: str) -> None:
    """Принудительно установить шрифт в rPr элемента."""
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f"<w:rFonts {nsdecls('w')} "
            f'w:ascii="{tag}" w:hAnsi="{tag}" '
            f'w:eastAsia="{tag}" w:cs="{tag}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(key), tag)


def _ppr(style_element) -> object:
    """Получить или создать pPr."""
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        style_element.append(pPr)
    return pPr


def _set_outline_level(style, level: int) -> None:
    pPr = _ppr(style.element)
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level}"/>')
        pPr.append(outline)
    else:
        outline.set(qn("w:val"), str(level))


def _add_page_break_before_style(style) -> None:
    pPr = _ppr(style.element)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pPr.append(parse_xml(f"<w:pageBreakBefore {nsdecls('w')}/>"))


def _set_cell_margins(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f"<w:tcMar {nsdecls('w')}>"
        '<w:top w:w="55" w:type="dxa"/>'
        '<w:left w:w="85" w:type="dxa"/>'
        '<w:bottom w:w="55" w:type="dxa"/>'
        '<w:right w:w="85" w:type="dxa"/>'
        "</w:tcMar>"
    )
    tcPr.append(tcMar)


def _set_cell_shading(cell, color: str = "D9E2F3") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shd)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f"<w:tblBorders {nsdecls('w')}>"
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def _set_table_width_pct(table, pct: int = 5000) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{pct}" w:type="pct"/>')
    tblPr.append(tblW)


def _set_header_footer_pPr(paragraph) -> object:
    """Получить или создать pPr для header/footer параграфа."""
    tag_ppr = qn("w:pPr")
    pPr = paragraph._p.find(tag_ppr)
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        paragraph._p.insert(0, pPr)
    return pPr


def _get_sect_pr(doc) -> object:
    """Получить sectPr."""
    return doc.sections[0]._sectPr


# ---------- стили ----------


def _get_or_create_style(doc, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def setup_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_BODY
    _rf(normal.element, FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # DPO-Body: body с красной строкой и justified
    body = _get_or_create_style(doc, "DPO-Body")
    body.base_style = normal
    body.font.name = FONT_NAME
    body.font.size = FONT_BODY
    pf = body.paragraph_format
    pf.first_line_indent = INDENT_BODY
    pf.line_spacing = LINE_SPACING_BODY
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _rf(body.element, FONT_NAME)

    # DPO-H1: 14pt bold caps, CENTER, outline=0, pageBreakBefore
    h1 = _get_or_create_style(doc, "DPO-H1")
    h1.font.name = FONT_NAME
    h1.font.size = FONT_BODY
    h1.font.bold = True
    h1.font.all_caps = True
    pf = h1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    _set_outline_level(h1, 0)
    _add_page_break_before_style(h1)
    _rf(h1.element, FONT_NAME)

    # DPO-H2: 14pt bold, LEFT, outline=1
    h2 = _get_or_create_style(doc, "DPO-H2")
    h2.font.name = FONT_NAME
    h2.font.size = FONT_BODY
    h2.font.bold = True
    pf = h2.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    _set_outline_level(h2, 1)
    _rf(h2.element, FONT_NAME)

    # DPO-H3: 14pt bold italic, LEFT, outline=2
    h3 = _get_or_create_style(doc, "DPO-H3")
    h3.font.name = FONT_NAME
    h3.font.size = FONT_BODY
    h3.font.bold = True
    h3.font.italic = True
    pf = h3.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = INDENT_BODY
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    _set_outline_level(h3, 2)
    _rf(h3.element, FONT_NAME)

    # DPO-Cell: для ячеек таблиц
    cell_style = _get_or_create_style(doc, "DPO-Cell")
    cell_style.font.name = FONT_NAME
    cell_style.font.size = FONT_TABLE
    pf = cell_style.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.line_spacing = LINE_SPACING_TABLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _rf(cell_style.element, FONT_NAME)


# ---------- страница ----------


def setup_page(doc) -> None:
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = MARGIN_LEFT
        sec.right_margin = MARGIN_RIGHT
        sec.top_margin = MARGIN_TOP
        sec.bottom_margin = MARGIN_BOTTOM


def set_document_language(doc, lang_ru: str = "ru-RU") -> None:
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = parse_xml(f"<w:docDefaults {nsdecls('w')}/>")
        styles_element.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = parse_xml(f"<w:rPrDefault {nsdecls('w')}/>")
        docDefaults.insert(0, rPrDefault)
    default_rPr = rPrDefault.find(qn("w:rPr"))
    if default_rPr is None:
        default_rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        rPrDefault.append(default_rPr)
    old = default_rPr.find(qn("w:lang"))
    if old is not None:
        default_rPr.remove(old)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), lang_ru)
    lang.set(qn("w:eastAsia"), lang_ru)
    lang.set(qn("w:bidi"), "ar-SA")
    default_rPr.append(lang)

    for name in ("Normal", "DPO-H1", "DPO-H2", "DPO-H3", "DPO-Body", "DPO-Cell"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
            style.element.insert(0, rPr)
        old = rPr.find(qn("w:lang"))
        if old is not None:
            rPr.remove(old)
        style_lang = OxmlElement("w:lang")
        style_lang.set(qn("w:val"), lang_ru)
        style_lang.set(qn("w:eastAsia"), lang_ru)
        style_lang.set(qn("w:bidi"), "ar-SA")
        rPr.append(style_lang)


def setup_page_numbers(doc) -> None:
    for sec in doc.sections:
        sectPr = sec._sectPr
        if sectPr.find(qn("w:titlePg")) is None:
            sectPr.append(parse_xml(f"<w:titlePg {nsdecls('w')}/>"))

        header = sec.header
        for p in list(header.paragraphs):
            p.clear()
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE \\* MERGEFORMAT"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(placeholder)
        r.append(fld_end)


# ---------- базовые помощники ----------


def _append_inline(p, text: str) -> None:
    """Обработка **bold** и *italic* в тексте."""
    if not text:
        return
    for segment in re.split(r"(\*\*.*?\*\*|\*[^*\n]+\*)", text):
        if not segment:
            continue
        bold = False
        italic = False
        if segment.startswith("**") and segment.endswith("**") and len(segment) >= 4:
            segment = segment[2:-2]
            bold = True
        elif segment.startswith("*") and segment.endswith("*") and len(segment) >= 2:
            segment = segment[1:-1]
            italic = True
        run = p.add_run(segment)
        run.font.name = FONT_NAME
        run.font.size = FONT_BODY
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        _rf(run._r, FONT_NAME)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_title_line(
    doc,
    text,
    *,
    bold=False,
    size=FONT_BODY,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=Pt(0),
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = LINE_SPACING_BODY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    if bold:
        run.font.bold = True
    _rf(run._r, FONT_NAME)
    return p


# ---------- титульник ----------


def add_dpo_title_page(doc, meta: dict) -> None:
    """Титульный лист ДПО-программы."""
    uni_name = meta.get("university", "Образовательная организация")
    license_no = meta.get("license_no", "")
    program_type = (
        "Повышение квалификации"
        if meta.get("program_type") == "pk"
        else "Профессиональная переподготовка"
    )

    _add_title_line(doc, uni_name, bold=True, size=Pt(14), space_after=Pt(24))
    if license_no:
        _add_title_line(
            doc,
            f"Лицензия на осуществление образовательной деятельности №{license_no}",
            size=Pt(12),
            space_after=Pt(36),
        )

    _add_title_line(doc, program_type, bold=True, size=Pt(16), space_after=Pt(12))
    _add_title_line(
        doc,
        f"\u00ab{meta.get('title', '\u2026')}\u00bb",
        bold=True,
        size=Pt(16),
        space_after=Pt(36),
    )

    code = meta.get("specialty_code", "")
    spec_name = meta.get("specialty_name", "")
    if code:
        _add_title_line(
            doc, f"Код специальности: {code} {spec_name}", space_after=Pt(12)
        )
    total_hrs = meta.get("total_hours", "—")
    form = meta.get("form", "—")
    duration = meta.get("duration_months", "—")
    audience = meta.get("audience", "—")
    city = meta.get("city", "г. Москва")
    year = meta.get("year", "2026")
    author = meta.get("author", "")

    _add_title_line(
        doc,
        f"Объём: {total_hrs} академических часов",
        space_after=Pt(12),
    )
    _add_title_line(
        doc,
        f"Форма обучения: {form}",
        space_after=Pt(12),
    )
    _add_title_line(
        doc,
        f"Срок реализации: {duration} мес.",
        space_after=Pt(12),
    )
    _add_title_line(
        doc,
        f"Целевая аудитория: {audience}",
        space_after=Pt(48),
    )

    if author:
        _add_title_line(doc, "Разработал(а):", space_after=Pt(6))
        _add_title_line(doc, author, space_after=Pt(48))

    _add_title_line(doc, city, space_after=Pt(0))
    _add_title_line(doc, str(year), space_after=Pt(0))

    # titlePg marker
    for sec in doc.sections:
        sectPr = sec._sectPr
        if sectPr.find(qn("w:titlePg")) is None:
            sectPr.append(parse_xml(f"<w:titlePg {nsdecls('w')}/>"))


# ---------- оглавление ----------


def add_toc(doc) -> None:
    _add_page_break(doc)
    _add_title_line(
        doc,
        "СОДЕРЖАНИЕ",
        bold=True,
        size=FONT_BODY,
        space_after=Pt(0),
    )
    p = doc.add_paragraph(style=doc.styles["DPO-Body"])
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_BODY
    _rf(run._r, FONT_NAME)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Для обновления оглавления нажмите F9 или ПКМ → Обновить поле."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


# ---------- таблицы ----------


def _looks_numeric(text: str) -> bool:
    t = (
        text.replace(" ", "")
        .replace(",", ".")
        .replace("\u2014", "")
        .replace("-", "")
        .strip()
    )
    if not t:
        return False
    return bool(re.match(r"^-?\d+(?:\.\d+)?%?$", t))


def _clean_md_inline(text: str) -> str:
    return re.sub(r"\*+", "", text)


def add_markdown_table(doc, md_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in md_lines:
        s = line.strip()
        if not s:
            continue
        if re.match(r"^\|[-:\s|]+\|?$", s):
            continue
        cells = [c.strip() for c in s.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    if len(rows) < 2:
        return

    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    _set_table_width_pct(table, 5000)

    # header row
    header_row = table.rows[0]
    header_row_trPr = header_row._tr.get_or_add_trPr()
    if header_row_trPr.find(qn("w:tblHeader")) is None:
        header_row_trPr.append(parse_xml(f"<w:tblHeader {nsdecls('w')}/>"))

    for ri, row_data in enumerate(rows):
        for ci, raw_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.style = doc.styles["DPO-Cell"]
            pf = p.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.line_spacing = LINE_SPACING_TABLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            txt = _clean_md_inline(raw_text)
            if ri == 0:
                _set_cell_shading(cell, "D9E2F3")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(txt)
                run.font.name = FONT_NAME
                run.font.size = FONT_TABLE
                run.font.bold = True
                _rf(run._r, FONT_NAME)
            else:
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                    if _looks_numeric(txt)
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                run = p.add_run(txt)
                run.font.name = FONT_NAME
                run.font.size = FONT_TABLE
                _rf(run._r, FONT_NAME)


def add_list_item(doc, text: str, *, ordered: bool = False, number: str = "") -> None:
    p = doc.add_paragraph(style=doc.styles["DPO-Body"])
    prefix = f"{number}. " if ordered else "\u2014 "
    run = p.add_run(prefix)
    run.font.name = FONT_NAME
    run.font.size = FONT_BODY
    _rf(run._r, FONT_NAME)
    _append_inline(p, text)


# ---------- ПАРСЕР ----------


def parse_dpo_markdown(doc, md_text: str, base_dir: Path) -> None:
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---" or re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # H1
        if re.match(r"^#\s+", line):
            title = re.sub(r"^#\s+", "", line).strip()
            p = doc.add_paragraph(style=doc.styles["DPO-H1"])
            _append_inline(p, title)
            # page break на H1
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
                p._p.insert(0, pPr)
            if pPr.find(qn("w:pageBreakBefore")) is None:
                pPr.append(parse_xml(f"<w:pageBreakBefore {nsdecls('w')}/>"))
            i += 1
            continue

        # H2
        if re.match(r"^##\s+", line):
            title = re.sub(r"^##\s+", "", line).strip()
            p = doc.add_paragraph(style=doc.styles["DPO-H2"])
            _append_inline(p, title)
            i += 1
            continue

        # H3
        if re.match(r"^###\s+", line):
            title = re.sub(r"^###\s+", "", line).strip()
            p = doc.add_paragraph(style=doc.styles["DPO-H3"])
            _append_inline(p, title)
            i += 1
            continue

        # Markdown table
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[-:\s|]+\|?$", lines[i + 1].strip())
        ):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        # Ordered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            add_list_item(doc, m_num.group(2), ordered=True, number=m_num.group(1))
            i += 1
            continue

        # Bullet list (только если начинается с — или - с пробелом, но не *)
        if re.match(r"^[-—–]\s+", stripped):
            text = re.sub(r"^[-—–]\s+", "", stripped)
            add_list_item(doc, text, ordered=False)
            i += 1
            continue

        # Обычный body — очищаем от артефактов markdown
        # Удаляем одиночные # в начале строки (не заголовки)
        if stripped.startswith("#"):
            stripped = re.sub(r"^#+\s*", "", stripped)
        # Удаляем ** и * для inline-форматирования
        stripped = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        stripped = re.sub(r"(?<!\w)\*(?!\w)", "", stripped)

        # Обычный body
        p = doc.add_paragraph(style=doc.styles["DPO-Body"])
        _append_inline(p, stripped)
        i += 1


# ---------- main ----------


def build_docx(
    md_paths: list[Path], meta: dict, output_path: Path, base_dir: Path | None = None
) -> None:
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    set_document_language(doc)

    # Титульник
    add_dpo_title_page(doc, meta)

    # Оглавление
    add_toc(doc)

    # Содержание тем
    for md_path in md_paths:
        if not md_path.exists():
            print(f"WARNING: not found: {md_path}", file=sys.stderr)
            continue
        md_text = md_path.read_text(encoding="utf-8")
        parse_dpo_markdown(doc, md_text, base_dir=base_dir or md_path.parent)

    setup_page_numbers(doc)
    doc.save(str(output_path))
    print(f"Saved: {output_path}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="md -> docx \u0440\u0435\u043d\u0434\u0435\u0440\u0435\u0440 \u0434\u043b\u044f \u0414\u041f\u041e-\u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c."
    )
    parser.add_argument(
        "--meta",
        required=True,
        help="JSON \u0441 \u043c\u0435\u0442\u0430\u0434\u0430\u043d\u043d\u044b\u043c\u0438 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u044b",
    )
    parser.add_argument(
        "--topics",
        nargs="+",
        help="\u0421\u043f\u0438\u0441\u043e\u043a markdown-\u0444\u0430\u0439\u043b\u043e\u0432 \u0442\u0435\u043c",
    )
    parser.add_argument(
        "--all",
        action="store_true",
        help="\u0410\u0432\u0442\u043e\u043c\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u0438 \u043d\u0430\u0439\u0442\u0438 \u0432\u0441\u0435 topic-*.md \u0432 \u043f\u0430\u043f\u043a\u0435 program/",
    )
    parser.add_argument(
        "--output",
        default="output.docx",
        help="\u0412\u044b\u0445\u043e\u0434\u043d\u043e\u0439 docx",
    )
    parser.add_argument(
        "--base-dir",
        help="\u0411\u0430\u0437\u043e\u0432\u0430\u044f \u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440\u0438\u044f \u0434\u043b\u044f \u043f\u043e\u0438\u0441\u043a\u0430 \u0444\u0430\u0439\u043b\u043e\u0432",
    )
    args = parser.parse_args(argv)

    base_dir = Path(args.base_dir) if args.base_dir else Path.cwd()

    # \u0417\u0430\u0433\u0440\u0443\u0437\u043a\u0430 meta.json
    meta: dict = {}
    meta_path = Path(args.meta)
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))

    # \u041e\u043f\u0440\u0435\u0434\u0435\u043b\u0435\u043d\u0438\u0435 \u0444\u0430\u0439\u043b\u043e\u0432 \u0442\u0435\u043c
    md_paths: list[Path] = []
    if args.topics:
        md_paths = [Path(p) for p in args.topics]
    elif args.all:
        program_dir = base_dir / "program"
        if program_dir.exists():
            md_paths = sorted(program_dir.glob("topic-*.md"))
        if not md_paths:
            md_paths = sorted(base_dir.glob("topic-*.md"))

    if not md_paths:
        print("ERROR: no topics specified. Use --topics or --all", file=sys.stderr)
        return 1

    output_path = Path(args.output)
    if not output_path.is_absolute():
        output_path = Path.cwd() / output_path

    build_docx(md_paths, meta, output_path, base_dir=base_dir)
    return 0


if __name__ == "__main__":
    sys.exit(main())
